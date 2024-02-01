import docx
import os


#opens the nea project document
path = "DOCUMENT PATH"
doc = docx.Document(path)

#reading the contents of the document
def read():
    para = doc.paragraphs
    print(len(para))
    n = 1
    while n < len(para):
        single_para = doc.paragraphs[n]
        for run in single_para.runs:
            print(run.text)
        n+=1

#adding new paragraph to document
def add_paragraph(text):
    doc.add_paragraph(text)
    doc.save(path)

#function to add an image to the document
def add_image(img):
    doc.add_picture(img,width=docx.shared.Inches(5), height=docx.shared.Inches(5))
    doc.save(path)

def open_excel():
    os.system("start EXCEL.EXE arc.xlsx")

